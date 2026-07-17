"""
api/utils/docx_generator.py
~~~~~~~~~~~~~~~~~~~~~~~~~~~
Utility to convert structured/markdown resume content into a .docx file.
"""

from io import BytesIO
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_resume_docx(tailored_resume_text: str) -> BytesIO:
    """
    Generates an in-memory .docx file from the tailored resume text.
    Approximates Jake's Resume Template format (Times New Roman, 
    small margins, clear headers).
    """
    document = Document()

    # Set up basic styles
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Adjust margins to fit more content
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    lines = tailored_resume_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # Name Header
            clean_text = line.lstrip('#').strip()
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_text)
            heading = document.add_paragraph()
            run = heading.add_run(clean_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(24)
            run.bold = True
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading.paragraph_format.space_after = Pt(2)
        elif line.startswith('Contact:'):
            # Contact Info
            clean_text = line.replace('Contact:', '').strip()
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_text)
            contact_p = document.add_paragraph(clean_text)
            contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_p.paragraph_format.space_after = Pt(12)
        elif line.startswith('## '):
            # Section Header
            clean_text = line.lstrip('#').strip().upper()
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_text)
            heading = document.add_paragraph()
            run = heading.add_run(clean_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            heading.paragraph_format.space_after = Pt(2)
            # Add a bottom border representation using a horizontal line
            border_p = document.add_paragraph()
            border_run = border_p.add_run('_' * 80)
            border_run.font.name = 'Times New Roman'
            border_run.font.size = Pt(10)
            border_p.paragraph_format.space_after = Pt(6)
        elif line.startswith('### '):
            # Role / Title Header
            clean_text = line.lstrip('#').strip()
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_text)
            role_p = document.add_paragraph()
            run = role_p.add_run(clean_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True
            role_p.paragraph_format.space_after = Pt(2)
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet point
            clean_line = line[2:].strip()
            clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_line)
            p = document.add_paragraph(clean_line, style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
        else:
            # Plain paragraph
            clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            p = document.add_paragraph(clean_line)
            p.paragraph_format.space_after = Pt(2)

    # Save to in-memory buffer
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    return file_stream
